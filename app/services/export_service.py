from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER


class ExportService:

    def generate_docx(self, resume) -> bytes:
        """Generate a professional .docx resume."""
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        info = resume.personal_info or {}

        # ---- HEADER: Name & Title ----
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(info.get("full_name", "Your Name").upper())
        name_run.bold = True
        name_run.font.size = Pt(22)
        name_run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        if info.get("title"):
            title_para = doc.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.add_run(info["title"])
            title_run.font.size = Pt(12)
            title_run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

        # Contact line
        contacts = []
        for field in ["email", "phone", "location", "website", "linkedin"]:
            if info.get(field):
                contacts.append(info[field])

        if contacts:
            contact_para = doc.add_paragraph()
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_run = contact_para.add_run(" | ".join(contacts))
            contact_run.font.size = Pt(9)
            contact_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Divider
        self._add_divider(doc)

        # ---- SUMMARY ----
        if resume.summary:
            self._add_section_heading(doc, "PROFESSIONAL SUMMARY")
            summary_para = doc.add_paragraph(resume.summary)
            summary_para.runs[0].font.size = Pt(10)

        # ---- EXPERIENCE ----
        if resume.experience:
            self._add_section_heading(doc, "EXPERIENCE")
            for exp in resume.experience:
                # Company & Dates row
                exp_para = doc.add_paragraph()
                company_run = exp_para.add_run(exp.get("company", ""))
                company_run.bold = True
                company_run.font.size = Pt(11)

                date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
                exp_para.add_run(f"\t{date_str}").font.size = Pt(9)

                # Position
                pos_para = doc.add_paragraph()
                pos_run = pos_para.add_run(exp.get("position", ""))
                pos_run.italic = True
                pos_run.font.size = Pt(10)
                pos_run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)

                if exp.get("location"):
                    pos_para.add_run(f" | {exp['location']}").font.size = Pt(9)

                # Bullets
                for bullet in exp.get("description", []):
                    if bullet.strip():
                        bullet_para = doc.add_paragraph(bullet, style="List Bullet")
                        bullet_para.runs[0].font.size = Pt(10)

        # ---- EDUCATION ----
        if resume.education:
            self._add_section_heading(doc, "EDUCATION")
            for edu in resume.education:
                edu_para = doc.add_paragraph()
                inst_run = edu_para.add_run(edu.get("institution", ""))
                inst_run.bold = True
                inst_run.font.size = Pt(11)

                date_str = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
                edu_para.add_run(f"\t{date_str}").font.size = Pt(9)

                deg_para = doc.add_paragraph()
                deg_str = edu.get("degree", "")
                if edu.get("field"):
                    deg_str += f" in {edu['field']}"
                deg_run = deg_para.add_run(deg_str)
                deg_run.italic = True
                deg_run.font.size = Pt(10)

                if edu.get("gpa"):
                    deg_para.add_run(f" | GPA: {edu['gpa']}").font.size = Pt(9)

        # ---- SKILLS ----
        if resume.skills:
            self._add_section_heading(doc, "SKILLS")
            for skill_group in resume.skills:
                skills_para = doc.add_paragraph()
                cat_run = skills_para.add_run(f"{skill_group.get('category', 'Skills')}: ")
                cat_run.bold = True
                cat_run.font.size = Pt(10)
                items_str = ", ".join(skill_group.get("items", []))
                skills_para.add_run(items_str).font.size = Pt(10)

        # ---- PROJECTS ----
        if resume.projects:
            self._add_section_heading(doc, "PROJECTS")
            for proj in resume.projects:
                proj_para = doc.add_paragraph()
                proj_run = proj_para.add_run(proj.get("name", ""))
                proj_run.bold = True
                proj_run.font.size = Pt(11)

                if proj.get("technologies"):
                    tech_str = " | " + ", ".join(proj["technologies"])
                    proj_para.add_run(tech_str).font.size = Pt(9)

                desc_para = doc.add_paragraph(proj.get("description", ""))
                desc_para.runs[0].font.size = Pt(10)

        # ---- CERTIFICATIONS ----
        if resume.certifications:
            self._add_section_heading(doc, "CERTIFICATIONS")
            for cert in resume.certifications:
                cert_para = doc.add_paragraph()
                cert_run = cert_para.add_run(cert.get("name", ""))
                cert_run.bold = True
                cert_run.font.size = Pt(10)
                cert_para.add_run(f" — {cert.get('issuer', '')}").font.size = Pt(10)
                if cert.get("date"):
                    cert_para.add_run(f" ({cert['date']})").font.size = Pt(9)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

    def _add_section_heading(self, doc, text: str):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x6C, 0x63, 0xFF)
        # Bottom border
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6C63FF')
        pBdr.append(bottom)
        pPr.append(pBdr)
        doc.add_paragraph()

    def _add_divider(self, doc):
        para = doc.add_paragraph()
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '6C63FF')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def generate_pdf(self, resume) -> bytes:
        """Generate a PDF resume using ReportLab."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=0.75 * inch,
            leftMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch
        )

        info = resume.personal_info or {}
        styles = getSampleStyleSheet()
        story = []

        # Custom styles
        name_style = ParagraphStyle(
            "NameStyle",
            parent=styles["Normal"],
            fontSize=22,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#1A1A2E"),
            alignment=TA_CENTER,
            spaceAfter=4
        )
        title_style = ParagraphStyle(
            "TitleStyle",
            parent=styles["Normal"],
            fontSize=12,
            textColor=colors.HexColor("#6C63FF"),
            alignment=TA_CENTER,
            spaceAfter=4
        )
        contact_style = ParagraphStyle(
            "ContactStyle",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#555555"),
            alignment=TA_CENTER,
            spaceAfter=8
        )
        section_style = ParagraphStyle(
            "SectionStyle",
            parent=styles["Normal"],
            fontSize=11,
            fontName="Helvetica-Bold",
            textColor=colors.HexColor("#6C63FF"),
            spaceBefore=10,
            spaceAfter=4
        )
        body_style = ParagraphStyle(
            "BodyStyle",
            parent=styles["Normal"],
            fontSize=10,
            spaceAfter=3
        )

        # Header
        name = info.get("full_name", "Your Name").upper()
        story.append(Paragraph(name, name_style))

        if info.get("title"):
            story.append(Paragraph(info["title"], title_style))

        contacts = [info[f] for f in ["email", "phone", "location"] if info.get(f)]
        if contacts:
            story.append(Paragraph(" | ".join(contacts), contact_style))

        story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#6C63FF")))
        story.append(Spacer(1, 8))

        # Summary
        if resume.summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
            story.append(Spacer(1, 4))
            story.append(Paragraph(resume.summary, body_style))

        # Experience
        if resume.experience:
            story.append(Paragraph("EXPERIENCE", section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
            story.append(Spacer(1, 4))
            for exp in resume.experience:
                date_str = f"{exp.get('start_date', '')} – {exp.get('end_date', 'Present')}"
                company_line = f"<b>{exp.get('company', '')}</b> &nbsp;&nbsp; {date_str}"
                story.append(Paragraph(company_line, body_style))
                pos = exp.get("position", "")
                loc = exp.get("location", "")
                pos_line = f"<i><font color='#6C63FF'>{pos}</font></i>"
                if loc:
                    pos_line += f" | {loc}"
                story.append(Paragraph(pos_line, body_style))
                for bullet in exp.get("description", []):
                    if bullet.strip():
                        story.append(Paragraph(f"• {bullet}", body_style))
                story.append(Spacer(1, 4))

        # Education
        if resume.education:
            story.append(Paragraph("EDUCATION", section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
            story.append(Spacer(1, 4))
            for edu in resume.education:
                date_str = f"{edu.get('start_date', '')} – {edu.get('end_date', '')}"
                story.append(Paragraph(f"<b>{edu.get('institution', '')}</b>  {date_str}", body_style))
                deg = edu.get("degree", "")
                if edu.get("field"):
                    deg += f" in {edu['field']}"
                story.append(Paragraph(f"<i>{deg}</i>", body_style))
                story.append(Spacer(1, 4))

        # Skills
        if resume.skills:
            story.append(Paragraph("SKILLS", section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC")))
            story.append(Spacer(1, 4))
            for sg in resume.skills:
                cat = sg.get("category", "Skills")
                items = ", ".join(sg.get("items", []))
                story.append(Paragraph(f"<b>{cat}:</b> {items}", body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer.getvalue()
